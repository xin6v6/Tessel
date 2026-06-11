#!/usr/bin/env python3
"""
gen_docx.py — 生成 .docx 文件
用法：python3 gen_docx.py '<JSON>'

JSON 格式：
{
  "output": "output.docx",          // 输出路径（必填）
  "title": "文档标题",               // 可选，文档大标题
  "sections": [                      // 内容段落列表
    { "heading": "一级标题", "level": 1 },
    { "text": "普通段落文字" },
    { "bullets": ["条目1", "条目2"] },
    { "table": { "headers": ["列A","列B"], "rows": [["1","2"],["3","4"]] } },
    { "code": "def foo():\n    pass" }
  ]
}
"""
import sys
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 输出路径必须在 FILE_AGENT_ROOT（默认 tmp/）内，防止路径遍历
ROOT = Path(os.environ.get("FILE_AGENT_ROOT", "tmp")).resolve()

def safe_output(raw: str) -> Path:
    p = (ROOT / raw).resolve() if not Path(raw).is_absolute() else Path(raw).resolve()
    if p != ROOT and not str(p).startswith(str(ROOT) + os.sep):
        print(f"路径越界：{raw} 不在允许目录 {ROOT} 内", file=sys.stderr)
        sys.exit(1)
    return p

def main():
    if len(sys.argv) < 2:
        print("用法: python3 gen_docx.py '<JSON>'", file=sys.stderr)
        sys.exit(1)

    data = json.loads(sys.argv[1])
    output_path = safe_output(data["output"])
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    title = data.get("title")
    if title:
        h = doc.add_heading(title, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in data.get("sections", []):
        if "heading" in section:
            doc.add_heading(section["heading"], level=section.get("level", 1))
        elif "text" in section:
            doc.add_paragraph(section["text"])
        elif "bullets" in section:
            for item in section["bullets"]:
                doc.add_paragraph(item, style="List Bullet")
        elif "numbered" in section:
            for item in section["numbered"]:
                doc.add_paragraph(item, style="List Number")
        elif "code" in section:
            p = doc.add_paragraph()
            run = p.add_run(section["code"])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif "table" in section:
            tbl = section["table"]
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            cols = max(len(headers), max((len(r) for r in rows), default=0))
            table = doc.add_table(rows=1 + len(rows), cols=cols)
            table.style = "Table Grid"
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    table.rows[ri + 1].cells[ci].text = str(cell)

    doc.save(str(output_path))
    print(str(output_path.resolve()))

if __name__ == "__main__":
    main()
